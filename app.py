import streamlit as st
import os
from google import genai
from PIL import Image
from docx import Document
import io
import socket
import qrcode
import re

API_KEY = "AQ.Ab8RN6JeD7GDknfM9jFK3SNq7eMlc0iKMp8pEAk9NZLgw17wzA"
client = genai.Client(api_key=API_KEY)

def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('10.255.255.255', 1))
        IP = s.getsockname()[0]
    except Exception:
        IP = '127.0.0.1'
    finally:
        s.close()
    return IP

def get_sort_key(file):
    name = os.path.splitext(file.name)[0]
    try:
        return int(name) 
    except ValueError:
        return name 

def clean_text_for_word(text):
    text = text.replace(r'\div', '÷')
    text = text.replace(r'\times', '×')
    text = text.replace(r'\frac{1}{2}', '1/2')
    text = text.replace(r'\frac{1}{3}', '1/3')
    text = text.replace(r'\frac{1}{4}', '1/4')
    text = text.replace(r'\frac{1}{5}', '1/5')
    text = text.replace(r'\frac{1}{6}', '1/6')
    text = text.replace(r'\frac{1}{7}', '1/7')
    text = text.replace(r'\frac{1}{8}', '1/8')
    text = text.replace(r'\frac{1}{9}', '1/9')
    
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    text = text.replace('**', '')
    text = text.replace('$', '')
    return text

def analyze_math_image(image):
    prompt = """
    이 이미지에 있는 수학 문제를 분석해서 아래 양식에 맞춰 완벽한 해설을 작성해줘.
    *주의*: 나누기 기호는 '\\div' 같은 코드를 쓰지 말고 반드시 일반 기호인 '÷'를 쓰고, 곱하기는 '×', 분수는 '1/2' 같은 형태로 직접 기호로 적어줘. LaTeX 코드를 절대 사용하지 마.
    각 문제의 구분은 반드시 '===='로 해줘.
    
    [Q] 문제 내용 (번호와 문제 텍스트)
    [CONCEPT] 이 문제를 푸는 데 필요한 핵심 개념이나 공식
    [STEP] 아이가 이해하기 쉬운 상세한 단계별 풀이 과정
    [TIP] 아빠의 친절한 조언 및 실수 방지 팁
    ====
    """
    response = client.models.generate_content(
        model='gemini-3.6-flash',
        contents=[image, prompt],
    )
    return response.text

def parse_math_blocks(analyzed_blocks_list):
    parsed_data = []
    for blocks in analyzed_blocks_list:
        for block in blocks.split('===='):
            if not block.strip():
                continue
            
            q_text, concept_text, step_text, tip_text = "", "", "", ""
            current_section = None
            
            for line in block.strip().split('\n'):
                line_str = line.strip()
                if line_str.startswith('[Q]'):
                    current_section = 'q'
                    q_text += line_str.replace('[Q]', '').strip() + " "
                elif line_str.startswith('[CONCEPT]'):
                    current_section = 'concept'
                    concept_text += line_str.replace('[CONCEPT]', '').strip() + " "
                elif line_str.startswith('[STEP]'):
                    current_section = 'step'
                    step_text += line_str.replace('[STEP]', '').strip() + "\n"
                elif line_str.startswith('[TIP]'):
                    current_section = 'tip'
                    tip_text += line_str.replace('[TIP]', '').strip() + " "
                else:
                    if current_section == 'q': q_text += line_str + " "
                    elif current_section == 'concept': concept_text += line_str + " "
                    elif current_section == 'step': step_text += line_str + "\n"
                    elif current_section == 'tip': tip_text += line_str + " "
            
            if not q_text and not step_text:
                q_text = block.strip()
                step_text = "상세 풀이를 생성했습니다."
            
            parsed_data.append({
                "q": q_text.strip() or "문제 내용",
                "concept": concept_text.strip() or "핵심 개념 정리",
                "step": step_text.strip() or "단계별 풀이 과정",
                "tip": tip_text.strip() or "꼼꼼하게 확인해보세요!"
            })
    return parsed_data

def create_math_word_document(parsed_data):
    doc = Document()
    doc.add_heading('🧮 아빠표 AI 수학 맞춤 해설서', 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '문제 (Question)'
    hdr_cells[1].text = '친절한 단계별 풀이 (Step-by-Step)'
    
    for item in parsed_data:
        row_cells = table.add_row().cells
        
        q_clean = clean_text_for_word(item['q'])
        concept_clean = clean_text_for_word(item['concept'])
        step_clean = clean_text_for_word(item['step'])
        tip_clean = clean_text_for_word(item['tip'])
        
        row_cells[0].text = q_clean
        row_cells[1].text = f"[핵심 개념]\n{concept_clean}\n\n[단계별 풀이]\n{step_clean}\n\n[아빠의 꿀팁]\n{tip_clean}"
                
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

st.set_page_config(layout="wide")

with st.sidebar:
    st.markdown("### 📱 스마트폰으로 문제 찍기")
    local_ip = get_local_ip()
    local_url = f"http://{local_ip}:8501"
    
    qr = qrcode.make(local_url)
    buf = io.BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)
    
    st.image(buf, caption="스마트폰 카메라로 스캔하세요")

st.title("🧮 욘다를 위한 아빠표 AI 수학 선생님")

if 'math_data' not in st.session_state:
    st.session_state.math_data = None
if 'math_file' not in st.session_state:
    st.session_state.math_file = None

uploaded_files = st.file_uploader("수학 문제집 페이지 사진을 업로드하거나 스마트폰으로 찍어 올리세요. (여러 장 가능)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_files:
    uploaded_files.sort(key=get_sort_key)
    st.info(f"총 {len(uploaded_files)}장의 수학 페이지가 업로드되었습니다. 분석을 시작합니다!")
    
    if st.button("수학 해설 및 워드 파일 생성"):
        all_analyzed_blocks = []
        
        for idx, file in enumerate(uploaded_files):
            with st.spinner(f"[{idx+1}/{len(uploaded_files)}] '{file.name}' 문제 풀이를 분석 중입니다..."):
                image = Image.open(file)
                analyzed_result = analyze_math_image(image)
                all_analyzed_blocks.append(analyzed_result)
        
        st.session_state.math_data = parse_math_blocks(all_analyzed_blocks)
        st.session_state.math_file = create_math_word_document(st.session_state.math_data)

if st.session_state.math_data:
    st.success("✅ 수학 해설서 생성이 완료되었습니다!")
    
    st.download_button(
        label="📥 맞춤형 수학 해설서(.docx) 다운로드",
        data=st.session_state.math_file,
        file_name="수학해설서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    st.subheader("👀 해설 미리보기")
    for item in st.session_state.math_data:
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**[문제]**\n{item['q']}")
        with col2:
            st.markdown(f"**[핵심 개념]** {item['concept']}")
            st.markdown(f"**[단계별 풀이]**\n{item['step']}")
            st.markdown(f"**[아빠의 꿀팁]** {item['tip']}")
        st.divider()
